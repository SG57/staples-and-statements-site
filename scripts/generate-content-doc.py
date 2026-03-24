"""
generate-content-doc.py — Create a Google Docs-friendly .docx for Danielle
===========================================================================
Reads src/i18n/en.json and generates docs/CONTENT-REVIEW.docx organized
by page with clear instructions for editing.
"""

import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).parent.parent
EN_JSON = ROOT / 'src' / 'i18n' / 'en.json'
OUT_DOCX = ROOT / 'docs' / 'CONTENT-REVIEW.docx'

# Brand colors
CHARCOAL = RGBColor(0x1A, 0x16, 0x13)
ACCENT = RGBColor(0xB8, 0x97, 0x7E)
MUTED = RGBColor(0x6B, 0x5E, 0x53)
LIGHT_GRAY = RGBColor(0x99, 0x99, 0x99)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = CHARCOAL
    return h


def add_instruction(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = LIGHT_GRAY
    run.font.italic = True
    return p


def add_field(doc, label, value, is_long=False):
    """Add a labeled text field that Danielle can edit."""
    # Label
    p = doc.add_paragraph()
    label_run = p.add_run(f'{label}:')
    label_run.font.size = Pt(9)
    label_run.font.bold = True
    label_run.font.color.rgb = ACCENT

    if isinstance(value, list):
        value = '\n'.join(str(v) for v in value)
    if not value or str(value).strip() == '':
        value = '(empty)'

    # Value - on new line for long text, same line for short
    if is_long or len(str(value)) > 100 or '\n' in str(value):
        vp = doc.add_paragraph()
        vr = vp.add_run(str(value).replace('\\n', '\n'))
        vr.font.size = Pt(11)
        vr.font.color.rgb = CHARCOAL
        # Add a thin line after long fields
        sep = doc.add_paragraph()
        sep_run = sep.add_run('─' * 60)
        sep_run.font.size = Pt(6)
        sep_run.font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)
    else:
        p.add_run('  ')
        vr = p.add_run(str(value))
        vr.font.size = Pt(11)
        vr.font.color.rgb = CHARCOAL


def flatten_section(data, prefix=''):
    """Flatten a nested dict into (key, value) pairs."""
    items = []
    for k, v in data.items():
        key = f'{prefix}.{k}' if prefix else k
        if isinstance(v, dict):
            items.extend(flatten_section(v, key))
        elif isinstance(v, list):
            for i, item in enumerate(v):
                if isinstance(item, dict):
                    items.extend(flatten_section(item, f'{key}[{i}]'))
                else:
                    items.append((f'{key}[{i}]', str(item)))
        else:
            items.append((key, str(v)))
    return items


# Page display config: (json_key, display_name, description, is_content_focus)
PAGES = [
    ('about', 'About Page', 'Your personal brand story — written in YOUR voice.', True),
    ('services_content', 'Service Descriptions', 'Each service\'s title, price, tagline, and full description.', True),
    ('home', 'Home Page', 'Landing page text — first impressions matter.', True),
    ('archetypes', 'Archetypes Page', 'The framework introduction for feminine and masculine archetypes.', True),
    ('archetypes_content', 'Archetype Descriptions', 'Each archetype\'s name, tagline, and full body text.', True),
    ('clients', 'Clients Page', 'Client experience section — social proof and process.', True),
    ('contact', 'Contact Page', 'Booking and contact information.', True),
    ('shop', 'Shop Page', 'Style guides and curated picks section.', True),
    ('services', 'Services Listing Page', 'Service page structure — tier labels, how-it-works steps, CTA.', True),
    ('guides', 'Style Guides', 'Digital product names, prices, and descriptions.', True),
    ('booking_confirmed', 'Booking Confirmation', 'What clients see after booking.', False),
]

# Structural sections (appendix)
STRUCTURAL = [
    ('nav', 'Navigation'),
    ('footer', 'Footer'),
    ('marquee', 'Scrolling Ticker'),
    ('ui', 'Shared UI Text (buttons, labels, badges)'),
    ('services_detail', 'Service Detail Page Template'),
    ('archetypes_detail', 'Archetype Detail Page Template'),
]


def generate():
    with open(EN_JSON, 'r', encoding='utf-8') as f:
        data = json.load(f)

    doc = Document()

    # Title page
    title = doc.add_heading('Staples & Statements', level=0)
    for run in title.runs:
        run.font.color.rgb = CHARCOAL
        run.font.size = Pt(28)

    subtitle = doc.add_heading('Website Content Review', level=1)
    for run in subtitle.runs:
        run.font.color.rgb = ACCENT
        run.font.size = Pt(18)

    doc.add_paragraph('')
    add_instruction(doc,
        'Hi Danielle! This document contains ALL the text on your website. '
        'You can edit any text directly in this document using Google Docs\' '
        '"Suggesting" mode (click the pencil icon → Suggesting). '
        'When you\'re done, just let Cord know and he\'ll update the site.\n\n'
        'The document is organized by page. Each section shows the current text '
        'with a label telling you what it is. Edit the text below each label — '
        'that\'s what shows on the site.\n\n'
        'TIP: Focus on the first sections (About, Services, Home) — those are '
        'the most important. The "Appendix" at the end has small structural text '
        'like button labels that you probably don\'t need to change.'
    )

    doc.add_page_break()

    # ── MAIN CONTENT SECTIONS ──
    add_heading(doc, 'Content Pages', level=1)
    add_instruction(doc, 'These are the main pages visitors see. Edit freely!')

    for json_key, display_name, description, is_focus in PAGES:
        section_data = data.get(json_key, {})
        if not section_data:
            continue

        doc.add_page_break()
        add_heading(doc, display_name, level=2)
        add_instruction(doc, description)
        doc.add_paragraph('')

        # For service/archetype content collections, show each item as a subsection
        if json_key in ('services_content', 'archetypes_content', 'guides'):
            for item_key, item_data in section_data.items():
                if isinstance(item_data, dict):
                    name = item_data.get('title', item_data.get('name', item_key))
                    add_heading(doc, str(name), level=3)

                    for field_key, field_val in item_data.items():
                        if field_key in ('slug', 'order', 'status', 'featured', 'gender', 'icon', 'color'):
                            continue  # skip technical fields
                        is_long = field_key in ('body', 'description') or len(str(field_val)) > 150
                        nice_label = field_key.replace('_', ' ').replace('perfectFor', 'Perfect For').replace('shortDescription', 'Short Description').replace('priceNote', 'Price Note').title()
                        add_field(doc, nice_label, field_val, is_long=is_long)
                else:
                    add_field(doc, item_key, item_data)
        else:
            # Regular page — flatten and display
            items = flatten_section(section_data)
            current_prefix = ''
            for key, value in items:
                # Add sub-headers for major sections
                parts = key.split('.')
                if len(parts) >= 2 and parts[0] != current_prefix:
                    current_prefix = parts[0]
                    nice_section = current_prefix.replace('_', ' ').title()
                    add_heading(doc, nice_section, level=3)

                # Nice label
                nice_key = key.split('.')[-1].replace('_', ' ').replace('[', ' #').replace(']', '').title()
                is_long = len(value) > 150 or '\n' in value
                add_field(doc, nice_key, value, is_long=is_long)

    # ── APPENDIX: STRUCTURAL UI ──
    doc.add_page_break()
    add_heading(doc, 'Appendix: Structural UI Text', level=1)
    add_instruction(doc,
        'These are small structural elements like button labels, navigation links, '
        'and footer text. You probably don\'t need to change these, but they\'re here '
        'if you want to review them.'
    )

    for json_key, display_name in STRUCTURAL:
        section_data = data.get(json_key, {})
        if not section_data:
            continue

        add_heading(doc, display_name, level=2)
        items = flatten_section(section_data)
        for key, value in items:
            nice_key = key.replace('_', ' ').replace('[', ' #').replace(']', '').title()
            add_field(doc, nice_key, value)

    # Save
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOCX))
    print(f'\n  Generated: {OUT_DOCX}')
    print(f'  Sections: {len(PAGES)} content + {len(STRUCTURAL)} structural')
    print(f'\n  Next steps:')
    print(f'  1. Upload {OUT_DOCX.name} to Google Drive')
    print(f'  2. Open with Google Docs')
    print(f'  3. Share with Danielle (Editor access)')
    print(f'  4. Tell her to use "Suggesting" mode')


if __name__ == '__main__':
    print('Staples & Statements — Content Review Doc Generator')
    generate()
    print()
